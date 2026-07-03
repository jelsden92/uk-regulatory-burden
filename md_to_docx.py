"""Minimal Markdown -> .docx converter for the coverage methodology note.

Handles the constructs actually used in docs/coverage_methodology_note.md: ATX headings
(#/##), bold (**...**), pipe tables, blockquotes (> ... rendered as a shaded
PROVISIONAL callout), and dash bullet lists. Not a general Markdown engine -- just
enough to render this note faithfully, preserving its provisional banners.

Usage: python md_to_docx.py [in.md] [out.docx]
"""
from __future__ import annotations
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_runs(paragraph, text):
    """Render inline `code` spans and **bold** within a paragraph. Backtick spans
    render monospace (backticks stripped); underscores inside them are left intact
    (identifiers like `en_core_web_trf` are NOT mangled)."""
    for ci, seg in enumerate(text.split('`')):
        if ci % 2 == 1:                       # inside a `code` span
            if seg:
                paragraph.add_run(seg).font.name = 'Consolas'
            continue
        for bi, part in enumerate(re.split(r'\*\*', seg)):
            if part == '':
                continue
            run = paragraph.add_run(part)
            run.bold = (bi % 2 == 1)


def shade(paragraph, hexcolor='FFF3CD'):
    """Apply a background shade to a paragraph (for the PROVISIONAL callouts)."""
    p = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    p.append(shd)


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else 'docs/coverage_methodology_note.md'
    out = sys.argv[2] if len(sys.argv) > 2 else 'coverage_methodology_note.docx'
    lines = open(src, encoding='utf-8').read().splitlines()

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        # blank
        if not line.strip():
            i += 1; continue

        # table block: consecutive lines starting with '|'
        if line.lstrip().startswith('|'):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                block.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip().strip('|').split('|')] for r in block]
            rows = [r for r in rows if not all(set(c) <= set('-: ') for c in r)]  # drop --- sep
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = 'Light Grid Accent 1'
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(t.rows[ri].cells):
                            cpar = t.rows[ri].cells[ci].paragraphs[0]
                            add_runs(cpar, cell)
                            if ri == 0:
                                for run in cpar.runs:
                                    run.bold = True
            continue

        # blockquote block: consecutive '> ' lines -> shaded callout
        if line.lstrip().startswith('>'):
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', lines[i])); i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.15)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            add_runs(para, ' '.join(b for b in buf if b.strip()))
            shade(para)
            continue

        # headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            doc.add_heading('', level=min(level, 4))  # placeholder for spacing
            h = doc.paragraphs[-1]
            add_runs(h, m.group(2))
            i += 1; continue

        # bullet list
        if re.match(r'^\s*[-*]\s+', line):
            para = doc.add_paragraph(style='List Bullet')
            add_runs(para, re.sub(r'^\s*[-*]\s+', '', line))
            i += 1; continue

        # whole-line italic wrapper (_..._) — e.g. the version/status line
        mital = re.fullmatch(r'_(.+)_', line.strip())
        if mital:
            para = doc.add_paragraph()
            add_runs(para, mital.group(1))
            for r in para.runs:
                r.italic = True
            i += 1; continue

        # ordinary paragraph
        para = doc.add_paragraph()
        add_runs(para, line)
        i += 1

    doc.save(out)
    print(f'Wrote {out}  ({len(doc.paragraphs)} paragraphs, '
          f'{len(doc.tables)} tables)')


if __name__ == '__main__':
    main()
